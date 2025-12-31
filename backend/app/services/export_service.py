"""
Export service for generating PRD documents.
"""
import logging
from datetime import datetime
from typing import Optional, Literal
from uuid import UUID
from io import BytesIO
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from backend.app.models.conversation import Conversation, Message
from backend.app.models.knowledge_base import KnowledgeBase
from backend.app.models.project import Project
from backend.app.services.gemini_service import GeminiService

# 导入导出相关的库
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown2
    MARKDOWN2_AVAILABLE = True
except ImportError:
    MARKDOWN2_AVAILABLE = False

logger = logging.getLogger(__name__)

ExportFormat = Literal['markdown', 'pdf', 'word', 'html']


class ExportService:
    """Service for exporting conversations as PRD documents."""
    
    def __init__(self, gemini_service: GeminiService):
        self.gemini_service = gemini_service
    
    async def export_conversation_to_markdown(
        self,
        db: AsyncSession,
        conversation_id: UUID,
        include_knowledge_base: bool = True,
        custom_template: Optional[str] = None
    ) -> tuple[str, str]:
        """
        Export a conversation as a Markdown PRD document.
        
        Args:
            db: Database session
            conversation_id: Conversation ID
            include_knowledge_base: Whether to include knowledge base
            custom_template: Custom template (optional)
            
        Returns:
            Tuple of (content, filename)
        """
        # Get conversation with messages
        conv_result = await db.execute(
            select(Conversation).where(Conversation.id == conversation_id)
        )
        conversation = conv_result.scalar_one_or_none()
        if not conversation:
            raise ValueError(f"Conversation {conversation_id} not found")
        
        # Get project
        project_result = await db.execute(
            select(Project).where(Project.id == conversation.project_id)
        )
        project = project_result.scalar_one()
        
        # Get messages
        messages_result = await db.execute(
            select(Message)
            .where(Message.conversation_id == conversation_id)
            .order_by(Message.sequence)
        )
        messages = messages_result.scalars().all()
        
        # Get knowledge base if requested
        knowledge_base = None
        if include_knowledge_base:
            kb_result = await db.execute(
                select(KnowledgeBase)
                .where(KnowledgeBase.project_id == conversation.project_id)
                .where(KnowledgeBase.status == "confirmed")
            )
            knowledge_base = kb_result.scalar_one_or_none()
        
        # Generate PRD using AI
        prd_content = await self._generate_prd_with_ai(
            project=project,
            conversation=conversation,
            messages=messages,
            knowledge_base=knowledge_base,
            custom_template=custom_template
        )
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = self._sanitize_filename(conversation.title or "PRD")
        filename = f"{safe_title}_{timestamp}.md"
        
        logger.info(f"Exported conversation {conversation_id} to PRD")
        
        return prd_content, filename
    
    async def _generate_prd_with_ai(
        self,
        project: Project,
        conversation: Conversation,
        messages: list,
        knowledge_base: Optional[KnowledgeBase],
        custom_template: Optional[str]
    ) -> str:
        """
        Use AI to generate a structured PRD from conversation.
        
        Args:
            project: Project object
            conversation: Conversation object
            messages: List of messages
            knowledge_base: Knowledge base (optional)
            custom_template: Custom template (optional)
            
        Returns:
            Generated PRD content in Markdown
        """
        # Build conversation summary
        conversation_text = self._format_conversation(messages)
        
        # Build knowledge base context
        kb_context = ""
        if knowledge_base:
            kb_context = self._format_knowledge_base(knowledge_base)
        
        # Build prompt for AI
        prompt = f"""基于以下对话内容，生成一份结构化的产品需求文档（PRD）。

# 项目信息
- 项目名称：{project.name}
- 项目描述：{project.description or '无'}

# 知识库
{kb_context if kb_context else '无项目知识库'}

# 对话内容
{conversation_text}

# 要求
请生成一份完整的PRD文档，包含以下部分：

1. **需求概述**
   - 需求背景
   - 需求目标
   - 目标用户

2. **功能需求**
   - 核心功能列表
   - 功能详细描述
   - 用户流程

3. **非功能需求**
   - 性能要求
   - 安全要求
   - 兼容性要求

4. **UI/UX 设计要求**
   - 界面设计规范
   - 交互设计要点
   - 视觉设计要求

5. **技术实现建议**
   - 技术架构
   - 接口设计
   - 数据模型

6. **验收标准**
   - 功能验收标准
   - 性能验收标准

7. **附录**
   - 相关文档
   - 参考资料

请使用 Markdown 格式，内容要详细、专业、可执行。
"""
        
        if custom_template:
            prompt = f"{custom_template}\n\n{prompt}"
        
        try:
            prd_content = await self.gemini_service.generate_text(
                prompt=prompt,
                temperature=0.3,  # Lower temperature for more structured output
                max_tokens=8000
            )
            
            # Add header
            header = f"""# {conversation.title or '产品需求文档'}

**项目名称**：{project.name}  
**文档版本**：1.0  
**创建日期**：{datetime.now().strftime('%Y-%m-%d')}  
**最后更新**：{datetime.now().strftime('%Y-%m-%d')}

---

"""
            
            return header + prd_content
            
        except Exception as e:
            logger.error(f"Error generating PRD with AI: {e}")
            # Fallback: return formatted conversation
            return self._generate_simple_prd(project, conversation, messages, knowledge_base)
    
    def _format_conversation(self, messages: list) -> str:
        """Format conversation messages as text."""
        lines = []
        for msg in messages:
            role = "PM" if msg.role == "user" else "AI助手"
            lines.append(f"**{role}**：{msg.content}\n")
        return "\n".join(lines)
    
    def _format_knowledge_base(self, kb: KnowledgeBase) -> str:
        """Format knowledge base as text."""
        data = kb.structured_data
        lines = []
        
        # System overview
        if data.get("system_overview"):
            overview = data["system_overview"]
            lines.append("## 系统概览")
            if overview.get("product_type"):
                lines.append(f"- 产品类型：{overview['product_type']}")
            if overview.get("core_modules"):
                lines.append(f"- 核心模块：{', '.join(overview['core_modules'])}")
            if overview.get("description"):
                lines.append(f"- 描述：{overview['description']}")
            lines.append("")
        
        # UI standards
        if data.get("ui_standards"):
            ui = data["ui_standards"]
            lines.append("## UI规范")
            if ui.get("primary_colors"):
                lines.append(f"- 主色调：{', '.join(ui['primary_colors'])}")
            if ui.get("component_library"):
                lines.append(f"- 组件库：{ui['component_library']}")
            if ui.get("layout_features"):
                lines.append(f"- 布局特征：{', '.join(ui['layout_features'])}")
            lines.append("")
        
        # Tech conventions
        if data.get("tech_conventions"):
            tech = data["tech_conventions"]
            lines.append("## 技术约定")
            if tech.get("naming_style"):
                lines.append(f"- 命名风格：{tech['naming_style']}")
            if tech.get("api_style"):
                lines.append(f"- API风格：{tech['api_style']}")
            lines.append("")
        
        return "\n".join(lines)
    
    def _generate_simple_prd(
        self,
        project: Project,
        conversation: Conversation,
        messages: list,
        knowledge_base: Optional[KnowledgeBase]
    ) -> str:
        """Generate a simple PRD without AI (fallback)."""
        lines = [
            f"# {conversation.title or '产品需求文档'}",
            "",
            f"**项目名称**：{project.name}",
            f"**文档版本**：1.0",
            f"**创建日期**：{datetime.now().strftime('%Y-%m-%d')}",
            "",
            "---",
            "",
            "## 需求对话记录",
            "",
        ]
        
        # Add conversation
        for msg in messages:
            role = "PM" if msg.role == "user" else "AI助手"
            lines.append(f"### {role}")
            lines.append(msg.content)
            lines.append("")
        
        # Add knowledge base
        if knowledge_base:
            lines.append("---")
            lines.append("")
            lines.append("## 项目知识库")
            lines.append("")
            lines.append(self._format_knowledge_base(knowledge_base))
        
        return "\n".join(lines)
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename by removing invalid characters."""
        # Remove invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')

        # Limit length
        if len(filename) > 50:
            filename = filename[:50]

        return filename.strip()

    async def export_conversation(
        self,
        db: AsyncSession,
        conversation_id: UUID,
        format: ExportFormat = 'markdown',
        include_knowledge_base: bool = True,
        custom_template: Optional[str] = None
    ) -> tuple[bytes | str, str, str]:
        """
        Export conversation to various formats.

        Args:
            db: Database session
            conversation_id: Conversation ID
            format: Export format (markdown, word, html, pdf)
            include_knowledge_base: Whether to include knowledge base
            custom_template: Custom template (optional)

        Returns:
            Tuple of (content, filename, content_type)
        """
        # First generate the markdown content
        markdown_content, base_filename = await self.export_conversation_to_markdown(
            db, conversation_id, include_knowledge_base, custom_template
        )

        # Convert to requested format
        if format == 'markdown':
            return markdown_content, base_filename, 'text/markdown'

        elif format == 'word':
            if not DOCX_AVAILABLE:
                raise RuntimeError("python-docx is not installed")

            word_content = self._markdown_to_word(markdown_content)
            word_filename = base_filename.replace('.md', '.docx')
            return word_content, word_filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        elif format == 'html':
            html_content = self._markdown_to_html(markdown_content)
            html_filename = base_filename.replace('.md', '.html')
            return html_content, html_filename, 'text/html'

        elif format == 'pdf':
            # PDF 导出需要额外的库，这里暂时返回 HTML
            # 实际生产环境可以使用 weasyprint 或其他库
            html_content = self._markdown_to_html(markdown_content)
            html_filename = base_filename.replace('.md', '.html')
            logger.warning("PDF export not yet implemented, returning HTML instead")
            return html_content, html_filename, 'text/html'

        else:
            raise ValueError(f"Unsupported export format: {format}")

    def _markdown_to_word(self, markdown_content: str) -> bytes:
        """
        Convert Markdown to Word document.

        Args:
            markdown_content: Markdown content

        Returns:
            Word document as bytes
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # 解析 Markdown 并添加到文档
        lines = markdown_content.split('\n')
        in_code_block = False
        code_lines = []

        for line in lines:
            # 检测代码块
            if line.strip().startswith('```'):
                if in_code_block:
                    # 结束代码块
                    if code_lines:
                        code_text = '\n'.join(code_lines)
                        p = doc.add_paragraph(code_text)
                        p.style = 'IntenseQuote'
                        code_lines = []
                    in_code_block = False
                else:
                    # 开始代码块
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            # 标题
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)

            # 分隔线
            elif line.strip() == '---':
                doc.add_paragraph('_' * 50)

            # 无序列表
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                p = doc.add_paragraph(text, style='List Bullet')

            # 有序列表
            elif line.strip() and line.strip()[0].isdigit() and '. ' in line:
                text = line.split('. ', 1)[1]
                p = doc.add_paragraph(text, style='List Number')

            # 引用
            elif line.startswith('>'):
                text = line[1:].strip()
                p = doc.add_paragraph(text)
                p.style = 'IntenseQuote'

            # 普通段落
            elif line.strip():
                # 处理加粗和斜体
                p = doc.add_paragraph(line)

            # 空行
            else:
                doc.add_paragraph()

        # 保存到字节流
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)

        return bio.getvalue()

    def _markdown_to_html(self, markdown_content: str) -> str:
        """
        Convert Markdown to HTML.

        Args:
            markdown_content: Markdown content

        Returns:
            HTML content
        """
        if MARKDOWN2_AVAILABLE:
            # 使用 markdown2 进行转换
            html_body = markdown2.markdown(
                markdown_content,
                extras=[
                    'tables',
                    'fenced-code-blocks',
                    'code-friendly',
                    'strike',
                    'task_list'
                ]
            )
        else:
            # 简单的 HTML 转换
            html_body = markdown_content.replace('\n', '<br>\n')

        # 包装在完整的 HTML 文档中
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>产品需求文档</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica', 'Arial', sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
            background: #fff;
        }}
        h1 {{
            font-size: 2.5em;
            margin-bottom: 0.5em;
            border-bottom: 3px solid #1890ff;
            padding-bottom: 0.3em;
        }}
        h2 {{
            font-size: 2em;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            border-bottom: 2px solid #e8e8e8;
            padding-bottom: 0.3em;
        }}
        h3 {{
            font-size: 1.5em;
            margin-top: 1.2em;
            margin-bottom: 0.5em;
        }}
        h4 {{
            font-size: 1.25em;
            margin-top: 1em;
            margin-bottom: 0.5em;
        }}
        p {{
            margin-bottom: 1em;
        }}
        code {{
            background: #f6f8fa;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'SFMono-Regular', Consolas, 'Liberation Mono', Menlo, monospace;
            font-size: 0.9em;
        }}
        pre {{
            background: #f6f8fa;
            padding: 16px;
            border-radius: 6px;
            overflow-x: auto;
            margin: 1em 0;
        }}
        pre code {{
            background: transparent;
            padding: 0;
        }}
        ul, ol {{
            margin-bottom: 1em;
            padding-left: 2em;
        }}
        li {{
            margin-bottom: 0.5em;
        }}
        blockquote {{
            border-left: 4px solid #dfe2e5;
            padding-left: 1em;
            color: #6a737d;
            margin: 1em 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        th, td {{
            border: 1px solid #dfe2e5;
            padding: 8px 12px;
            text-align: left;
        }}
        th {{
            background: #f6f8fa;
            font-weight: 600;
        }}
        hr {{
            border: none;
            border-top: 2px solid #e8e8e8;
            margin: 2em 0;
        }}
        a {{
            color: #1890ff;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        @media print {{
            body {{
                max-width: 100%;
                padding: 20px;
            }}
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

        return html

